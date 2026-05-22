from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


class UserUploadMarkdownConverter:
    def _normalize_text(self, text: str) -> str:
        lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
        cleaned: list[str] = []
        blank_streak = 0
        for line in lines:
            if not line:
                blank_streak += 1
                if blank_streak <= 1:
                    cleaned.append("")
                continue
            blank_streak = 0
            cleaned.append(line)
        return "\n".join(cleaned).strip()

    def _convert_pdf(self, input_file: Path) -> str:
        reader = PdfReader(str(input_file))
        pages: list[str] = []
        for page_index, page in enumerate(reader.pages, start=1):
            text = self._normalize_text(page.extract_text() or "")
            if text:
                pages.append(f"<!-- page: {page_index} -->\n\n{text}")
        return "\n\n".join(pages).strip()

    def _convert_docx(self, input_file: Path) -> str:
        document = DocxDocument(str(input_file))
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            text = self._normalize_text(paragraph.text)
            if text:
                blocks.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [self._normalize_text(cell.text) for cell in row.cells]
                row_text = " | ".join(cell for cell in cells if cell)
                if row_text:
                    blocks.append(row_text)
        return "\n\n".join(blocks).strip()

    def convert_to_markdown(self, input_path: str, output_path: str) -> str:
        input_file = Path(input_path)
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        suffix = input_file.suffix.lower()
        if suffix == ".pdf":
            markdown_text = self._convert_pdf(input_file)
        elif suffix == ".docx":
            markdown_text = self._convert_docx(input_file)
        else:
            raise ValueError(f"Unsupported user upload file type: {suffix}")

        if not markdown_text:
            raise ValueError("Could not extract text from uploaded document.")
        if not markdown_text.endswith("\n"):
            markdown_text += "\n"
        output_file.write_text(markdown_text, encoding="utf-8")
        return str(output_file)
