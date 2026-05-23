import asyncio
import json
import shutil
import traceback
from io import BytesIO
from pathlib import Path

from docx import Document
from starlette.datastructures import UploadFile

from app.db.mongodb import get_database, get_mongo_client
from app.rag.vectorstore.chroma_store import ChromaVectorStore
from app.schemas.chat_schema import ChatRequest
from app.services.chat_service import ChatService
from app.services.document_service import DocumentService
from app.services.session_service import SessionService
from app.services.user_service import UserService
from app.workers.ingestion_worker import IngestionWorker


PREFIX = "manual_e2e_rag"
EMAIL = f"{PREFIX}@example.test"
PASSWORD = "Password123!"
LOG_PATH = Path("manual_e2e_rag_check.log")


def log_step(message: str) -> None:
    with LOG_PATH.open("a", encoding="utf-8") as f:
        f.write(message + "\n")


async def cleanup(document_ids: list[str] | None = None) -> None:
    document_ids = document_ids or []
    db = get_database()
    user = await db.users.find_one({"email": EMAIL})
    user_id = user["_id"] if user else None
    session_ids: list[str] = []
    if user_id:
        session_ids = [item["_id"] async for item in db.sessions.find({"owner_user_id": user_id})]
        docs = [item async for item in db.documents.find({"owner_user_id": user_id})]
        document_ids.extend([doc["_id"] for doc in docs])
        for doc in docs:
            for path_key in ("raw_storage_path", "markdown_storage_path"):
                path_value = doc.get(path_key)
                if path_value:
                    remove_path_and_empty_parents(path_value)

    vector_store = ChromaVectorStore()
    for doc_id in set(document_ids):
        try:
            vector_store.delete_by_document_id(doc_id)
        except Exception:
            pass
        await db.chunks.delete_many({"document_id": doc_id})
        await db.ingestion_jobs.delete_many({"document_id": doc_id})
        await db.documents.delete_many({"_id": doc_id})

    if user_id:
        await db.messages.delete_many({"owner_user_id": user_id})
        await db.retrieval_logs.delete_many({"user_id": user_id})
        await db.sessions.delete_many({"owner_user_id": user_id})
        await db.users.delete_many({"_id": user_id})
    for session_id in session_ids:
        await db.messages.delete_many({"session_id": session_id})


def remove_path_and_empty_parents(path_value: str) -> None:
    path = Path(path_value)
    if path.exists():
        if path.is_dir():
            shutil.rmtree(path)
        else:
            path.unlink()
    parent = path.parent
    while parent.exists() and parent.name not in {"user_upload", "raw", "markdown", "storage"}:
        try:
            parent.rmdir()
        except OSError:
            break
        parent = parent.parent


def build_docx_upload() -> UploadFile:
    doc = Document()
    doc.add_heading("Ho so Alpha test RAG", level=1)
    doc.add_paragraph("Tai lieu nay la file user upload dung de kiem thu RAG trong session.")
    doc.add_paragraph("Le phi noi bo cua ho so Alpha la 75.000 dong, nop tai quay tiep nhan.")
    doc.add_paragraph("Thoi han xu ly noi bo cua ho so Alpha la 05 ngay lam viec ke tu khi nhan du ho so hop le.")
    doc.add_paragraph("Nguoi nop can mang giay de nghi va ban sao giay to phap ly ca nhan.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return UploadFile(file=buffer, filename="manual_alpha_e2e.docx")


async def run_check() -> dict:
    created_document_id = None
    if LOG_PATH.exists():
        LOG_PATH.unlink()
    log_step("start cleanup")
    await cleanup()
    try:
        log_step("register user")
        user, token = await UserService().register_user("Manual E2E RAG User", EMAIL, PASSWORD)
        log_step(f"create session user={user.id}")
        session = await SessionService().create_session(user.id, title="Manual E2E RAG Session")
        log_step(f"upload document session={session.id}")
        upload_result = await DocumentService().upload_user_document(
            build_docx_upload(),
            owner_user_id=user.id,
            session_id=session.id,
        )
        document = upload_result["document"]
        created_document_id = document.id
        job_id = upload_result["job_id"]

        log_step(f"claim ingestion job={job_id}")
        worker = IngestionWorker(poll_interval_seconds=0)
        job = await worker.ingestion_job_repository.claim_next_queued_job()
        if job is None:
            raise RuntimeError("No queued ingestion job was found")
        log_step(f"process ingestion job={job['_id']}")
        await worker._process_job(job)

        log_step("load ingestion result")
        db = get_database()
        doc_after = await db.documents.find_one({"_id": created_document_id})
        job_after = await db.ingestion_jobs.find_one({"_id": job_id})
        chunk_docs = [chunk async for chunk in db.chunks.find({"document_id": created_document_id})]

        question = "Trong file manual_alpha_e2e.docx vừa upload, lệ phí và thời hạn xử lý là bao nhiêu?"
        log_step("chat question")
        result = await ChatService().ask_question(
            ChatRequest(question=question, session_id=session.id, scope="auto"),
            user_id=user.id,
        )
        log_step("load messages")
        updated_session = await db.sessions.find_one({"_id": session.id})
        messages = [message async for message in db.messages.find({"session_id": session.id}).sort("created_at", 1)]

        return {
            "user": {"id": user.id, "email": user.email, "token_created": bool(token)},
            "session": {"id": session.id, "title": session.title},
            "upload": {
                "document_id": created_document_id,
                "filename": document.filename,
                "initial_status": document.status,
                "job_id": job_id,
            },
            "ingestion": {
                "document_status": doc_after.get("status") if doc_after else None,
                "job_status": job_after.get("status") if job_after else None,
                "job_step": job_after.get("current_step") if job_after else None,
                "chunk_count_document": doc_after.get("chunk_count") if doc_after else None,
                "chunk_count_db": len(chunk_docs),
                "markdown_storage_path": doc_after.get("markdown_storage_path") if doc_after else None,
            },
            "chat": {
                "question": question,
                "intent_resolution": result.get("intent_resolution"),
                "scope": result.get("scope"),
                "scope_resolution": result.get("scope_resolution"),
                "retrieved_chunk_ids": [
                    item.get("metadata", {}).get("chunk_id", item.get("id")) for item in result.get("raw_contexts", [])
                ],
                "sources": result.get("sources"),
                "answer_preview": result.get("answer", "")[:1500],
            },
            "session_after_chat": {
                "message_count": len(messages),
                "conversation_state": updated_session.get("conversation_state") if updated_session else None,
            },
        }
    finally:
        log_step("final cleanup")
        await cleanup([created_document_id] if created_document_id else [])


async def main() -> None:
    try:
        output = await run_check()
        print(json.dumps(output, ensure_ascii=False, indent=2, default=str))
    except Exception:
        traceback.print_exc()
        raise
    finally:
        get_mongo_client().close()
        get_mongo_client.cache_clear()


if __name__ == "__main__":
    asyncio.run(main())
